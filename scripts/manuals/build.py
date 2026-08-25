"""
Four user manuals, as Word documents.

WRITTEN FOR THE PERSON WHO HAS THE ROLE, not for somebody reading about the
software. That decides everything: a student's manual never mentions a
"section-subject", an admin's says what a decision COSTS rather than which
button makes it, and none of them contain the word "endpoint".

The structure of each is the same on purpose — Getting started, then one
chapter per screen in the order the sidebar lists it, then the things people
get wrong. Somebody who has read one can find their way around another, which
matters in an office where the same person is sometimes both.

Every screen listed here is one the role can actually reach: the lists come
from navigation.ts rather than from memory.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# The repository root, not docs/manuals. These are documents the Institute
# HANDS TO PEOPLE — a QA engineer, a new teacher — and burying them three
# directories down makes them something only a developer ever finds.
#
# Derived from this file's own location rather than written out, so a clone
# anywhere still writes to the right place.
OUT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))

NAVY = RGBColor(0x1A, 0x3C, 0x5E)
AMBER = RGBColor(0xC1, 0x7D, 0x11)
GREY = RGBColor(0x55, 0x5F, 0x6B)


# ───────────────────────────────────────────────────────────── helpers ──────
def styled(doc):
    """One look for all four, set once on the base styles."""
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(8)
    n.paragraph_format.line_spacing = 1.15

    for name, size, colour, before in [
        ("Heading 1", 20, NAVY, 22),
        ("Heading 2", 15, NAVY, 18),
        ("Heading 3", 12, AMBER, 12),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = colour
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(6)
    return doc


def cover(doc, role, subtitle, audience):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Prepreneurship")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Learning Management System")
    r.font.size = Pt(13)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(role)
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = AMBER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

    for _ in range(6):
        doc.add_paragraph()
    w = doc.add_paragraph()
    w.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = w.add_run(audience)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GREY
    doc.add_page_break()


def para(doc, text, bold=False, italic=False, grey=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if grey:
        r.font.color.rgb = GREY
        r.font.size = Pt(10)
    return p


def bullets(doc, items):
    for i in items:
        doc.add_paragraph(i, style="List Bullet")


def steps(doc, items):
    for i in items:
        doc.add_paragraph(i, style="List Number")


def callout(doc, label, text):
    """The thing people get wrong, set apart so it is not read as prose."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    p = c.paragraphs[0]
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = AMBER
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10)
    doc.add_paragraph()


def screen(doc, name, what_for, how, note=None):
    doc.add_heading(name, level=2)
    para(doc, what_for)
    if how:
        doc.add_heading("How to use it", level=3)
        steps(doc, how)
    if note:
        callout(doc, "Watch out:", note)


def signing_in(doc, who):
    doc.add_heading("Getting started", level=1)
    doc.add_heading("Signing in", level=2)
    steps(doc, [
        "Open the address the Institute gave you in any web browser.",
        "Enter your email address and password, then choose Sign in.",
        "The first time you sign in you will be asked to set your own password. "
        "Choose something only you know — the one you were given has been seen by the office.",
    ])
    callout(doc, "If you cannot sign in:",
            "Three wrong passwords in a row locks the account for a short while. It unlocks by "
            "itself; you do not need to do anything. If you have forgotten your password, ask the "
            f"office to reset it — {who}")

    doc.add_heading("Finding your way around", level=2)
    para(doc, "Everything is reached from the list down the left-hand side. On a phone, tap the "
              "menu button at the top to open it.")
    bullets(doc, [
        "Press Ctrl and K together to search for any screen by name.",
        "The bell at the top right shows anything new addressed to you.",
        "Your name at the bottom left is where you change your password or sign out.",
    ])
    doc.add_heading("The blue panel at the top of each screen", level=2)
    para(doc, "Most screens open with a short panel explaining what the screen is for, in four "
              "numbered steps, with the one thing people commonly get wrong underneath. Choose "
              "Hide once you no longer need it and it collapses to a single line — it stays there, "
              "so you can open it again months later.")
    doc.add_page_break()


def closing(doc, extra=None):
    doc.add_heading("If something looks wrong", level=1)
    bullets(doc, [
        "Reload the page first. Most oddities are a screen that has been open a long while.",
        "Check the date and the term at the top of the screen — a figure that looks wrong is "
        "usually the right figure for a different term.",
        "Note what you were doing and what you expected, and tell the office. If an error message "
        "showed a reference beginning ERR-, include it: it identifies the exact request in the log.",
    ])
    if extra:
        bullets(doc, extra)
    doc.add_paragraph()
    para(doc, "This manual describes the System as installed at the Prepreneurship Institute. "
              "Screens change as the Institute changes; the blue panel on each screen is always "
              "current, because it lives in the software rather than in this document.",
         grey=True)


os.makedirs(OUT, exist_ok=True)
print(f"writing to {OUT}")
