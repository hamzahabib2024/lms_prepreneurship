"""The teacher's manual. Written for somebody whose job is teaching and for
whom this software is overhead — so it is organised by what they are trying to
do, not by what the software contains."""

import os
from docx import Document
from build import (OUT, styled, cover, para, bullets, steps, callout, table, screen,
                   signing_in, closing)

doc = styled(Document())

cover(doc, "Teacher Guide",
      "Your classes, your marking, and your recordings",
      "For teaching staff at the Prepreneurship Institute")

doc.add_heading("What is in this guide", level=1)
table(doc, ["Section", "What it covers"], [
    ["Getting started", "Signing in, and what you can reach"],
    ["Attendance", "Taking the register, and who is at risk"],
    ["Setting work", "Assignments, spoken briefs, and marking guides"],
    ["Marking", "One submission at a time, and releasing marks"],
    ["Quizzes", "Building one, and marking written answers"],
    ["Recordings", "Connecting a Drive folder and publishing lectures"],
    ["Content", "Modules, lessons and course material"],
    ["Communicating", "Announcements and discussion"],
    ["Reports", "Getting the numbers out"],
    ["What you can and cannot see", "The limits of your role"],
])
doc.add_page_break()

signing_in(doc, "they can set you a new one, and you will choose your own when you next sign in.")

doc.add_heading("What you can reach", level=1)
para(doc, "You see the classes you are assigned to and nothing else. That is deliberate: it is "
          "not a matter of trust but of keeping one teacher's records out of another's screen. "
          "If a class you teach is missing, you have not been assigned to it — ask the office.")
callout(doc, "The most common first-day problem:",
        "A section with no teacher assigned is invisible to every teacher, so nobody takes the "
        "register. If your timetable looks empty, this is almost always why.")

doc.add_heading("Your screens, one by one", level=1)

screen(doc, "Attendance",
       "Taking the register, one class at a time. Everybody starts as present, so you only mark "
       "the exceptions — which is faster and less error-prone than marking thirty people present.",
       ["Choose the class. Today's are listed first.",
        "Mark anybody absent, late or excused. Leave everybody else alone.",
        "Choose Save. Nothing is recorded until you do.",
        "You can correct a register afterwards. The change is recorded with your name."],
       "Late is not absent, and the difference matters — late counts partially towards "
       "attendance. A student arriving after the class started is marked late, measured from "
       "when the class actually began, not when it was scheduled.")

doc.add_heading("Who is at risk", level=3)
para(doc, "The System warns a student by itself once their attendance drops below the "
          "requirement. You do not have to watch for it. It waits until enough classes have been "
          "held to mean anything — missing one of two classes is 50% attendance and says nothing.")

screen(doc, "Setting work",
       "An assignment is the task, when it is due, what happens if it is late, and how it is "
       "marked. Nothing reaches a student until you publish it, so you can prepare a whole term "
       "in advance.",
       ["Choose New assignment from the Marking screen.",
        "Give it a title and say what they have to do. You can also record a spoken brief — "
        "for design or language work, forty seconds of explanation beats four paragraphs.",
        "Set when it opens, when it is due, and what happens to late work.",
        "Say what they may hand in: typed answers, files, or a spoken answer.",
        "Choose a marking guide if you want the marks broken down.",
        "Save it as a draft, or publish it when you are ready."],
       "A grace period is worth setting. Every term somebody's upload finishes at 23:01, and "
       "marking that late punishes a slow connection rather than leaving it late.")

doc.add_heading("Marking guides", level=3)
para(doc, "A marking guide is a list of what you are looking for, with marks against each part. "
          "Write it once and every piece of work is judged the same way — and a student can see "
          "why they got what they got.")
steps(doc, [
    "Open Marking guides and choose New.",
    "Add one line per thing you are marking: idea, craft, presentation.",
    "Put marks against each. The running total is shown as you type.",
    "Choose it when you set an assignment.",
])
callout(doc, "It must add up:",
        "The marks in the guide have to total what the assignment is out of. The running total "
        "tells you where you are — a guide that does not add up cannot be used.")

screen(doc, "Marking",
       "The Marking screen is your queue: every class, every assignment, and how many "
       "submissions are still waiting. Open one to mark the work in it.",
       ["Choose an assignment. The number beside it is how many are still to mark.",
        "Choose Mark to open the first one that needs doing.",
        "The work, the marking guide and the mark box are on one screen. Enter the mark and any "
        "feedback, then Save.",
        "Saving takes you to the next one still to mark. The arrow keys move between them and "
        "Escape returns to the list.",
        "When you have finished the whole class, choose Release."],
       "Enter the RAW mark — what the work is worth. Any late penalty is worked out for you. If "
       "you deduct it yourself it is applied twice.")

doc.add_heading("Releasing marks", level=3)
para(doc, "Nothing you enter is visible to any student until you release the assignment. Then "
          "the whole class sees theirs at once.")
bullets(doc, [
    "This is why you can mark over several days without students comparing notes early.",
    "You can change a mark after releasing it, but the System asks for a reason and records it.",
    "Internal notes are never shown to a student. Feedback is.",
])

screen(doc, "Quizzes",
       "A quiz is its rules plus its questions. Multiple choice, true or false, short answer and "
       "numeric are supported, and the first two mark themselves.",
       ["Choose New quiz and set the rules: how long, how many attempts, whether the order is "
        "shuffled.",
        "Add questions, or pull them from a question bank.",
        "Decide when results appear — immediately, when the quiz closes, or only when you say.",
        "Publish it.",
        "Written answers appear on the quiz marking screen, grouped by question."],
       "Showing results immediately tells a student the answers. If the rest of the class has "
       "not sat it yet, choose 'when the quiz closes' instead.")

doc.add_heading("Marking written answers", level=3)
para(doc, "Only answers that need a person are listed. They are grouped BY QUESTION rather than "
          "by student, because marking the same question across the whole class in one pass "
          "gives a far more consistent standard than working through one student at a time.")

screen(doc, "Recordings",
       "Class recordings arrive in a Google Drive folder and appear here by themselves. You do "
       "not upload them twice.",
       ["Open Courses and choose the class.",
        "Connect the Drive folder its recordings land in. You do this once.",
        "Anything put in that folder is picked up and listed here.",
        "Check a recording plays before the class does.",
        "New recordings arrive unpublished — publish the ones you want students to see."],
       "Changing a class's folder clears the recordings from the old one and pulls in the new "
       "folder's. Nothing is destroyed: point it back and they return, with each student's "
       "watch progress intact.")

screen(doc, "Content",
       "Course material, in three levels: a course holds modules, a module holds lessons, and a "
       "lesson holds what students open.",
       ["Make a module — a block of the course.",
        "Add lessons to it. One lesson is one sitting.",
        "Attach notes, files and the recording.",
        "Publish it. Nothing is visible to a student until you do."],
       "Publishing a lesson inside an unpublished module shows nobody anything. The module has "
       "to be published too.")

screen(doc, "Announcements and Discussion",
       "Announcements reach a whole class or one subject. Discussion is where students ask "
       "questions and everybody can read the answer.",
       ["Choose the audience: one of your subjects, or a whole section.",
        "Write a clear title — many people read only that.",
        "Set an expiry if it is about an event, so it takes itself down.",
        "Post it. It reaches everybody in that audience immediately."],
       "You can announce to your own classes. Announcing to every teacher or to the whole "
       "Institute is the office's to do — you will be refused, and that is why.")

screen(doc, "Reports",
       "Counted from the Institute's records at the moment you ask, never from a stored summary.",
       ["Choose the report and narrow it to your class and term.",
        "Read it on screen before sending it anywhere.",
        "Export it as a spreadsheet if somebody outside needs it."],
       "An export stops being true the moment somebody's attendance changes. Put the date on it "
       "before you send it to anyone.")

doc.add_heading("What you can and cannot see", level=1)
table(doc, ["You can", "You cannot"], [
    ["See students in the classes you teach", "See students in classes you do not teach"],
    ["Mark work and release marks", "Change a student's enrolment or roll number"],
    ["Take and correct registers", "See any student's fees or payments"],
    ["Post to your own classes", "Post to the whole Institute or to all teachers"],
    ["Read institute policy on the Settings screen", "Change institute policy"],
    ["Issue nothing", "Issue or revoke certificates — that is the office"],
])
para(doc, "None of this is about trust. A teacher does not need the fee ledger to teach, and a "
          "system that hands out reach nobody needs is one where a mistake goes further than it "
          "should.", grey=True)

doc.add_page_break()
closing(doc, [
    "If a class is missing from your screens, you have not been assigned to it. The office can "
    "assign you in a moment.",
])

path = os.path.join(OUT, "Prepreneurship LMS - Teacher Guide.docx")
doc.save(path)
print("wrote", os.path.basename(path))
