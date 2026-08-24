"""The student's manual. The only one written for somebody who did not choose
to use this software — they enrolled on a course and this came with it."""

import os
from docx import Document
from build import (OUT, styled, cover, para, bullets, steps, callout, table, screen,
                   signing_in, closing)

doc = styled(Document())

cover(doc, "Student Guide",
      "Everything you can do, and how",
      "For students enrolled at the Prepreneurship Institute")

# ─────────────────────────────────────────────────────────────── contents ──
doc.add_heading("What is in this guide", level=1)
table(doc, ["Section", "What it covers"], [
    ["Getting started", "Signing in, finding your way around"],
    ["Dashboard", "What is due and what needs your attention"],
    ["My subjects", "Your courses, recordings, assignments and quizzes"],
    ["Handing work in", "Files, typed answers, and recording your voice"],
    ["Quizzes", "Sitting one, and seeing your result"],
    ["Timetable", "When your classes are and how to join"],
    ["Fees", "What you owe, and telling us you have paid"],
    ["My certificates", "Getting one, and proving it is genuine"],
    ["Announcements & Discussion", "Notices, and asking questions"],
    ["Your account", "Password, and what the Institute can see"],
])
doc.add_page_break()

signing_in(doc, "they can set you a new one, and you will choose your own when you next sign in.")

# ──────────────────────────────────────────────────────────── the screens ──
doc.add_heading("Your screens, one by one", level=1)

screen(doc, "Dashboard",
       "The first thing you see. It gathers what needs doing: work due soon, quizzes not yet "
       "attempted, your attendance, and how far through each course you are.",
       ["Read the top row first — anything overdue is shown in amber or red.",
        "Choose any item to go straight to it.",
        "The progress figure is how much of the course you have completed, counting lectures "
        "watched, work submitted, marks released and attendance."],
       "Your progress figure is not your grade. It is how much of the course you have DONE. You "
       "can have full marks on everything you have handed in and still be at 40% because half "
       "the term is still to come.")

screen(doc, "My subjects",
       "One card for each subject you are enrolled on. Open one to find its recordings, its "
       "assignments, its quizzes and your own progress in it.",
       ["Choose a subject to open it.",
        "Recordings are listed under the lesson they belong to. Choose one to watch it.",
        "The player remembers where you stopped — reopen it and it starts from there.",
        "Your watching is counted towards your progress, so let a lecture play rather than "
        "skipping to the end."],
       "Skipping to the end of a recording does not count it as watched. The System records "
       "which PARTS you have played, not how far along the bar you dragged it.")

screen(doc, "Handing work in",
       "Assignments appear in the subject they belong to. Each one says what is wanted, what it "
       "is out of, when it is due, and what happens if it is late.",
       ["Open the assignment and read the instructions. If your teacher recorded a spoken brief, "
        "choose Listen to hear it.",
        "Type your answer, attach files, or record your voice — whichever the assignment allows.",
        "Check the list of what you are about to send. You can remove anything before submitting.",
        "Choose Submit. You will see it confirmed straight away."],
       "Uploading a file is NOT submitting. The file sits in a list until you choose Submit. "
       "Every term somebody uploads their work an hour before the deadline, closes the page, and "
       "is marked as not having handed anything in.")

doc.add_heading("Recording your answer", level=3)
para(doc, "Where the assignment allows it, you can speak your answer instead of typing it — "
          "useful for language work, or explaining a design decision.")
steps(doc, [
    "Choose Record. Your browser will ask permission to use the microphone; choose Allow.",
    "Speak. The timer shows how long you have been recording and the limit.",
    "Choose Stop. You can play it back, and Record again if you want another go.",
    "It is added to your list of files like anything else. Then choose Submit.",
])
callout(doc, "If nothing records:",
        "The browser blocked the microphone. Look for the padlock in the address bar, allow the "
        "microphone for this site, and press Record again. On a phone, check no other app is "
        "using the microphone.")

screen(doc, "Quizzes",
       "Quizzes appear in the subject they belong to, with how long you get and how many "
       "attempts you are allowed.",
       ["Open the quiz and read the rules at the top before you start — the timer begins when "
        "you choose Start.",
        "Answer the questions. Your answers are saved as you go.",
        "Choose Submit when you are done, or it submits itself when the time runs out.",
        "Multiple-choice and numeric answers are marked immediately. Written answers wait for "
        "your teacher."],
       "Do not start a quiz to 'have a look'. Starting it begins the timer and uses one of your "
       "attempts, even if you close the page.")

screen(doc, "Timetable",
       "Your classes, with the room or the join link, the teacher, and the time.",
       ["Today's classes are at the top.",
        "For an online class, a Join button appears shortly before it starts.",
        "If a class has been moved or cancelled you will see it here and in your notices."],
       None)

screen(doc, "Fees",
       "What you have been charged, what you have paid, and what is left. Every payment the "
       "Institute has confirmed has a receipt you can open.",
       ["The top of the screen shows the balance outstanding.",
        "Each instalment shows its own due date.",
        "Choose Submit a fee payment to tell the Institute you have paid.",
        "Attach a clear photograph of the deposit slip or the transfer confirmation.",
        "The payment shows as awaiting confirmation until the office has checked it against "
        "the bank, then as paid, and a receipt appears."],
       "Telling us you have paid is not the same as us confirming it. Only a confirmed payment "
       "reduces your balance. Send the slip as soon as you have paid rather than on the due "
       "date, so there is time to check it.")

screen(doc, "My certificates",
       "Certificates the Institute has issued you, and the number that proves each one is real.",
       ["A certificate appears here once the Institute has issued it.",
        "Download it to print or attach to an application.",
        "Anybody — an employer, another institution — can check it is genuine by entering its "
        "number on the Institute's public page. They do not need an account."],
       "A certificate is issued when the requirements are genuinely met — attendance, work "
       "submitted, and marks. If you believe you have met them and no certificate has appeared, "
       "ask the office rather than waiting.")

screen(doc, "Announcements and Discussion",
       "Announcements are notices from the Institute or your teachers. Discussion is where you "
       "ask questions about a subject and get answers other students can read too.",
       ["Announcements addressed to you appear on this screen and in the bell at the top right.",
        "In Discussion, choose the subject, then ask your question.",
        "Read what has already been asked before posting — your question may already be "
        "answered."],
       None)

# ──────────────────────────────────────────────────────────── the account ──
doc.add_heading("Your account and your information", level=1)
doc.add_heading("Changing your password", level=2)
steps(doc, [
    "Choose your name at the bottom left.",
    "Choose Change password.",
    "Enter your current password, then the new one twice.",
])
doc.add_heading("What the Institute can see", level=2)
para(doc, "Being clear about this is easier than being asked:")
table(doc, ["The Institute can see", "The Institute cannot see"], [
    ["Your attendance, marks and submitted work", "Anything you type and do not submit"],
    ["Which recordings you have watched, and how much", "Your password — nobody can see it, ever"],
    ["Your fees and payments", "Your messages to anybody outside this System"],
    ["When you signed in, and from roughly where", "Any other website you visit"],
])
para(doc, "Your teacher sees this for the subjects they teach you. The office sees it for the "
          "Institute. Other students see none of it.")

doc.add_page_break()
closing(doc, [
    "If a recording will not play, try once more before reporting it — most failures are the "
    "connection rather than the file.",
])

path = os.path.join(OUT, "Prepreneurship LMS - Student Guide.docx")
doc.save(path)
print("wrote", os.path.basename(path))
