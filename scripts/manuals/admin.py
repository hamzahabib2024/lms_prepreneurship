"""The administrator's manual — the office. Organised around DECISIONS rather
than screens, because almost everything an admin does affects a real person's
record and several of them cannot be undone."""

import os
from docx import Document
from build import (OUT, styled, cover, para, bullets, steps, callout, table, screen,
                   signing_in, closing)

doc = styled(Document())

cover(doc, "Administrator Guide",
      "Running the Institute day to day",
      "For office and administrative staff")

doc.add_heading("What is in this guide", level=1)
table(doc, ["Batch", "What it covers"], [
    ["Getting started", "Signing in, and the shape of your role"],
    ["Admissions", "Turning an application into a student"],
    ["People", "Accounts, roles and suspensions"],
    ["Courses & fees", "What is taught and what it costs"],
    ["Batches & structure", "Classes, terms and intakes"],
    ["Online classes", "The meeting link a class joins every week"],
    ["Deleting things", "Removing what was created by mistake"],
    ["Fees & payments", "Charges, confirming payments, receipts"],
    ["Certificates", "Issuing, verifying, revoking"],
    ["Announcements", "Who a notice reaches, including the public page"],
    ["The public page", "Editing what visitors see"],
    ["Import & bulk changes", "Many records at once"],
    ["Reports & audit", "Numbers, and who changed what"],
    ["What only a Super Admin can do", "Where your role stops"],
])
doc.add_page_break()

signing_in(doc, "another administrator can reset it for you.")

doc.add_heading("The shape of your role", level=1)
para(doc, "You can see and change almost everything about students, courses and money. Two "
          "things you cannot do, and both are deliberate:")
bullets(doc, [
    "You can READ institute policy on the Settings screen but not change it. Those values "
    "decide when a student is warned and what a certificate requires — changing them is a "
    "governance decision, reserved to a Super Admin.",
    "You cannot read the security log, restore a backup, or grant roles at the highest level.",
])
callout(doc, "Anything touching money asks again:",
        "Opening the fee ledger or recording a payment asks you to confirm your password, even "
        "though you are already signed in. That is deliberate — it means an unattended screen "
        "cannot be used to move money.")

doc.add_heading("Your screens, one by one", level=1)

screen(doc, "Admissions",
       "Everything people have applied for, newest first. Each is a decision waiting to be made. "
       "Approving creates the student, their registration number and their sign-in, all at once.",
       ["Open the application and read it: their details, the course, and the payment slip.",
        "Check the slip — does the amount match the fee, and is the name the same? This is the "
        "step that catches most mistakes.",
        "Approve, decline, or ask for more information.",
        "They are told either way. An approval sends their sign-in details; a rejection says why."],
       "If something is merely missing, ask for it rather than declining. 'Needs more "
       "information' keeps their place in the queue; declining makes them start over.")

screen(doc, "People",
       "Everybody with a sign-in, and what each is allowed to reach.",
       ["Search by name or email.",
        "Give a role. The role decides what they see.",
        "Suspend rather than delete when somebody leaves — suspending stops them signing in "
        "without destroying anything they did.",
        "Every change is recorded with your name."],
       "Give the narrowest role that lets somebody do their job. Widening it later takes a "
       "moment; explaining why a part-time tutor could read the fee ledger does not.")

screen(doc, "Courses & fees",
       "What the Institute teaches and what it costs. This is where the public page gets its "
       "list from.",
       ["Create the course: name, code, and how long it runs.",
        "Add its subjects — the parts it is taught in.",
        "Set the fee: the full price, and the instalments if you allow them.",
        "Publish the price. Until you do, nobody is quoted anything."],
       "The instalment lines have to add up to the total. The running total under each block "
       "tells you where you are, so you are not re-adding twelve numbers on the day you publish.")

screen(doc, "Batches and Structure",
       "A batch is one group of students taught together — a class. Structure is the calendar "
       "the whole Institute runs on: terms and intakes.",
       ["In Structure, create the term and its intakes, with the dates.",
        "In Batches, create the class: its course, its name, its shift and its term.",
        "Assign a teacher."],
       "A batch with no teacher is invisible to every teacher, so nobody takes the register. "
       "It is the single most common thing to forget.")

screen(doc, "The meeting link for an online class",
       "Where a class is taught online, the subject inside the batch carries the link to its "
       "meeting room — one link, used every week, and permanent until somebody changes it. "
       "Teachers set their own; you can set any, which matters when a teacher is away or a "
       "class is covered.",
       ["Open Batches, choose the batch, and open its subjects.",
        "Each subject has its own teacher and its own Meeting link column. The link belongs to "
        "the subject INSIDE the class, not to the course — two batches of the same course meet "
        "in two different rooms at two different times.",
        "Choose the link cell — it reads 'not set' until there is one — and paste the whole "
        "address, including the https:// at the front.",
        "Add a note if students need one: a passcode, or when to join.",
        "Come back and choose 'Change the link' at any point. Clearing the box removes it."],
       "It must begin with https:// and is refused otherwise. You do not need it ready before "
       "the course opens — set it whenever you have it, and every student sees it immediately.")

# ─────────────────────────────────────────────────── deleting things ──
screen(doc, "Deleting a course, an intake or a subject",
       "Setting up a term means creating a lot of records quickly, and some of them are wrong: "
       "an intake named for the wrong session, a subject added to the wrong class, a batch "
       "created twice. You can delete any of them — while nothing depends on them.",
       ["Subjects: Courses & fees, then Delete on the subject's card.",
        "Intakes: Structure, then Delete on the intake's row.",
        "Batches: the Batches screen, then Delete on the batch's row.",
        "A subject on one class: open that batch's subjects and choose Remove on its row.",
        "Every one asks you to confirm first, and says what it is about to do."],
       "The System REFUSES to delete anything that has been used, and tells you exactly what is "
       "in the way — 'this intake still has 1 batch (AAB-A-MALE)'. Work down the list: take the "
       "subjects off the batch, then delete the batch, then the intake.")

doc.add_heading("Delete or archive?", level=3)
para(doc, "They are not two words for the same thing, and choosing wrong is the one mistake here "
          "that matters:")
table(doc, ["Archive", "Delete"], [
    ["For a class that FINISHED", "For a record that should never have existed"],
    ["Keeps attendance, marks and fees", "Only possible when there are none"],
    ["Leaves the active lists", "Leaves every list and dropdown"],
    ["Always available", "Refused the moment anything depends on it"],
])
para(doc, "A batch that has held one register is the only record that a student was marked "
          "present on a Tuesday in March. Archive it. The student who needs that record is "
          "usually applying for a job two years later, and by then nobody remembers the tidy-up.",
     grey=True)

screen(doc, "Fees and payments",
       "What every student owes, what they have paid, and the receipts.",
       ["Charges are raised from the published fee when a student enrols.",
        "A student can tell you they have paid and attach a slip; it appears on Payment "
        "verification.",
        "Check it against the bank, then confirm or reject it.",
        "Confirming reduces their balance and issues a receipt."],
       "A payment a student has told you about is not a payment. Only confirming it reduces "
       "what they owe. If you confirm without checking the bank, the ledger is wrong and the "
       "student has a receipt saying otherwise.")

screen(doc, "Certificates",
       "A certificate is a public claim by the Institute, so it is issued only when the "
       "requirements are genuinely met.",
       ["Check the student qualifies — the System tells you if attendance, work or marks are short.",
        "Issue it. It gets a number of its own, printed on the document.",
        "Anybody can verify that number on the public page without an account.",
        "Revoke only if you must; a revoked certificate stays on record and reads as revoked."],
       "What the certificate says is copied onto it at the moment you issue it. Renaming a "
       "course next year does not rewrite a certificate already in somebody's hands.")

screen(doc, "Announcements",
       "Notices, and who they reach. There are six audiences and the difference matters.",
       ["Choose the audience — the line under the box says exactly who will see it.",
        "Write a clear title. Many people read only that.",
        "Set an expiry for anything about an event.",
        "Post it."],
       None)
table(doc, ["Audience", "Who sees it"], [
    ["One of my subjects", "The students in that subject, and its teacher"],
    ["Everyone at the Institute", "Everybody with an account — students and staff"],
    ["Teachers only", "Teaching staff. Students never see it"],
    ["All staff", "Teachers and the office. Students never see it"],
    ["The public page only", "Visitors on the website, and NOBODY with an account"],
])
callout(doc, "The one worth re-reading:",
        "'The public page only' reaches no inbox at all. It is for an open day or an admissions "
        "deadline — people who have not applied yet. 'Everyone at the Institute' with the public "
        "box ticked reaches both, which puts an advertisement in front of students already "
        "enrolled.")

screen(doc, "The public page",
       "What a visitor sees before they have an account: the headline, the videos, the "
       "photographs, the claims and the closing band. All of it editable, live the moment you save.",
       ["Open Public page and work down the numbered batches.",
        "Show the preview to see the real page beside the form.",
        "Photographs upload directly — use pictures of the Institute rather than stock images.",
        "Videos take the ordinary share link from YouTube, TikTok, Facebook or Instagram.",
        "Save. It is live immediately."],
       "The programme list is NOT edited here. It comes from your real records, so a course you "
       "close stops being advertised the same afternoon. That is deliberate: a page that "
       "advertises a course that no longer exists takes applications for it.")

screen(doc, "Import and bulk changes",
       "For many records at once — a cohort joining from elsewhere, or moving a group between "
       "batches.",
       ["Download the template and fill it in, one row per student.",
        "Upload it. Nothing is created yet.",
        "Read the preview — it shows exactly what would happen, row by row.",
        "Fix anything flagged, upload again, and only then import."],
       "Always read the preview. It is the only chance to catch a wrong column before it becomes "
       "two hundred students on the wrong course. If a count is larger than you expected, your "
       "filter is wrong — not the System.")

screen(doc, "Reports and Audit",
       "Reports are counted at the moment you ask. The audit log is a permanent record of every "
       "change: who, what, when, and what it was before.",
       ["Choose a report, narrow it, read it on screen, then export if needed.",
        "In Audit, search by person, by kind of change, or by date.",
        "If somebody was acting as another user, the entry names both."],
       "Nobody can edit or delete the audit log — including you, and including a Super Admin. "
       "That is what makes it worth having.")

doc.add_heading("What only a Super Admin can do", level=1)
table(doc, ["Task", "Why it is not yours"], [
    ["Change institute policy (Settings)", "Those values decide when a student is warned and what "
     "a certificate requires"],
    ["Read the security log", "It names who has been attacked and from where — as useful for "
     "investigating a colleague as for defending one"],
    ["Take and restore backups", "A restore undoes every change since it was taken, for everybody"],
    ["Grant and revoke roles at the highest level", "The one power that could grant itself more"],
    ["Act as another user", "Produces a session indistinguishable from the real person's"],
])

doc.add_page_break()
closing(doc, [
    "If a figure looks wrong, check the term first. Most 'wrong' numbers are the right number "
    "for a different term.",
])

path = os.path.join(OUT, "Prepreneurship LMS - Administrator Guide.docx")
doc.save(path)
print("wrote", os.path.basename(path))
