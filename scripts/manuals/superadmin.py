"""The Super Admin's manual.

The only one of the four written mostly about RESTRAINT. Everything unique to
this role is either irreversible, invisible to the person it affects, or both —
so the manual's job is less "how do I" than "what happens if I do".
"""

import os
from docx import Document
from build import (OUT, styled, cover, para, bullets, steps, callout, table, screen,
                   signing_in, closing)

doc = styled(Document())

cover(doc, "Super Administrator Guide",
      "Governance, security and the things that cannot be undone",
      "For the Institute's system owner")

doc.add_heading("Read this first", level=1)
para(doc, "You can do everything an administrator can do, and this guide does not repeat the "
          "Administrator Guide — read that one for admissions, fees, certificates and the rest. "
          "This covers only what is yours alone.")
para(doc, "Almost everything in this guide is either irreversible, invisible to the person it "
          "affects, or both. None of it is needed on an ordinary day. If you find yourself using "
          "these screens weekly, something else is wrong.", bold=True)

doc.add_heading("What is in this guide", level=1)
table(doc, ["Section", "What it covers"], [
    ["Settings", "Institute policy, and what changing it does"],
    ["Roles and permissions", "Who can do what, and step-up"],
    ["Security", "Sign-ins, lockouts and sessions"],
    ["Backups and restoring", "The most destructive thing here"],
    ["Acting as another user", "What it is for and what it refuses"],
    ["Integrations", "Google, email and WhatsApp"],
    ["The audit log", "Why even you cannot edit it"],
    ["Deployment", "What breaks with more than one server"],
    ["A checklist", "Before and after going live"],
])
doc.add_page_break()

signing_in(doc, "there must be a second Super Admin who can reset it — see the checklist.")

doc.add_heading("Settings — institute policy", level=1)
para(doc, "These values decide when a student is warned, what counts as complete, what a "
          "certificate requires, and how progress is weighted. An administrator can read them; "
          "only you can change them.")
steps(doc, [
    "Search for the setting — the search matches its description, so you can look for what it does.",
    "Read what changing it does. Every setting says so in words.",
    "Change it and Save. The default is always shown beside it.",
])
callout(doc, "Changes are not retroactive:",
        "Lowering the attendance threshold does not un-warn students already warned. Re-weighting "
        "progress does not rewrite a certificate already issued. Anything already decided stays "
        "decided; these values inform decisions still to come.")
callout(doc, "If a change appears to do nothing:",
        "Look at where the value came from — shown on every row. A course or section may carry "
        "its own override, and the more specific one wins. This is the single most common reason "
        "somebody concludes the feature is broken.")

doc.add_heading("Roles and permissions", level=1)
para(doc, "Four roles: Super Admin, Admin, Teacher, Student. An Admin can additionally be granted "
          "individual sub-permissions — financial reporting, bulk operations, certificate issuing, "
          "managing other admins.")
bullets(doc, [
    "Grant the narrowest role that lets somebody do their job.",
    "Every grant is recorded against your name.",
    "Sensitive actions require step-up: you confirm your password again even though you are "
    "signed in. That is what stops an unattended screen being used.",
])
callout(doc, "Keep two Super Admins:",
        "A Super Admin cannot be impersonated and cannot have their password reset by an Admin. "
        "With only one, losing that password means losing governance of the System.")

doc.add_heading("Security", level=1)
para(doc, "Failed sign-ins, locked accounts, and anything resembling somebody trying to get in. "
          "Only you can read this — an Admin cannot, because it names who has been attacked and "
          "from where, which is as useful for investigating a colleague as for defending one.")
steps(doc, [
    "Look at what is flagged.",
    "Check the account — usually somebody has forgotten their password.",
    "Unlock or reset if needed. A lockout also clears itself with time.",
    "End sessions to sign somebody out everywhere, on every device, at once.",
])
callout(doc, "Look for the pattern:",
        "Most of what appears here is people mistyping their own password. Many accounts from "
        "one place, or one account from many places, is a different thing. Judge the shape "
        "before assuming an attack.")

doc.add_heading("Backups and restoring", level=1)
para(doc, "Copies of everything the Institute holds. Check them; test them; use them rarely.")
steps(doc, [
    "Check when the last one ran and whether it worked.",
    "Set how often they are taken and how long they are kept.",
    "Test a restore into a spare environment — a backup nobody has ever restored is a backup "
    "nobody knows works.",
])
callout(doc, "Restoring is the most destructive action in this System:",
        "It returns everything to the moment the backup was taken. Every change since — every "
        "mark, register, payment and enrolment, for everybody — is lost, not only the mistake you "
        "are fixing. It requires step-up, and it should require a conversation.")
para(doc, "The audit log is never restored over. It is append-only, and a restore that could "
          "rewrite the record of what happened would defeat the point of having one.", grey=True)

doc.add_heading("Acting as another user", level=1)
para(doc, "You can act as another user to see what they see. It exists because 'the page is "
          "broken for me' is otherwise unanswerable. It is the most dangerous capability here, "
          "and it refuses four things outright:")
table(doc, ["It refuses", "Why"], [
    ["Acting as another Super Admin", "Two people could each act as the other and no log would "
     "say who decided anything"],
    ["Acting as somebody while already acting as somebody", "The record holds one name, so a "
     "chain could not be traced"],
    ["Changing a password while acting", "That is account theft, not support — the owner is "
     "locked out and you can then sign in directly"],
    ["Continuing indefinitely", "The session expires by arithmetic and cannot be renewed"],
])
para(doc, "Every action taken while acting as somebody records both names.", grey=True)

doc.add_heading("Integrations", level=1)
screen(doc, "Google — Drive, Meet and Calendar",
       "Drive holds the recordings, Meet creates the class links, Calendar carries the timetable.",
       ["Paste the service-account credentials once.",
        "Test the connection — it says straight away whether it worked.",
        "Share each recordings folder with the service account as an Editor, not a Viewer.",
        "A folder must be on a Shared Drive, or uploads are refused however it is shared — a "
        "service account has no storage quota of its own."],
       "'The System can read the folder but cannot add to it' means it was shared as a Viewer. "
       "That is the most common Drive problem and the screen will say so.")

screen(doc, "WhatsApp",
       "Messages to students and parents, through the Meta Cloud API.",
       ["Set the access token and phone number id from the Meta Business account.",
        "Create a message template in WhatsApp Manager with TWO body parameters: {{1}} the title "
        "and {{2}} the message.",
        "Wait for Meta to approve it, then set its name and language code.",
        "The Integrations screen tells you which of the three is still missing."],
       "The template is not optional. Meta only accepts pre-approved templates outside a 24-hour "
       "reply window, and students never message the Institute first — so without one, nothing "
       "sends. Until all three are set, messages are simulated and readable on the Integrations "
       "screen.")

screen(doc, "Email",
       "Everything the Institute sends by mail: credentials, receipts, admission outcomes.",
       ["Set the SMTP host, port, user and password.",
        "For Gmail, the password is a 16-character App Password, not the account password.",
        "Send a test message and confirm it arrives before any student depends on it."],
       "A service shown as simulated is not sending anything. Safe for trying things out, and it "
       "must be switched on before real students rely on it.")

doc.add_heading("The audit log", level=1)
para(doc, "Every change: who, what, when, and what it was before. Nobody can edit or delete it — "
          "not an Admin, not you. The database itself refuses.")
para(doc, "That is deliberate. A record the most powerful person can rewrite is not a record. It "
          "exists so that 'who changed this mark' has an answer, and it is a poor tool for "
          "watching staff — it says what changed, not what anybody was doing.", grey=True)

doc.add_heading("Deployment — what breaks with more than one server", level=1)
para(doc, "If the Institute ever runs more than one copy of the API behind a load balancer, some "
          "things stop working in ways nothing reports. The System now checks at startup and says "
          "so in the log, worst first. Set API_INSTANCES so it knows how many there are.")
table(doc, ["What", "What goes wrong", "Fix"], [
    ["Local file storage", "A payment slip uploaded to one server does not exist on another. The "
     "record survives and the file does not", "Set DOCUMENT_STORAGE and LECTURE_STORAGE to Google "
     "Drive or object storage"],
    ["No shared cache", "Rate limits are counted per server, so three servers allow three times "
     "the password attempts before locking out", "Set REDIS_URL"],
    ["Proxy header unset", "Every request looks like it came from the proxy, so the whole "
     "Institute is rate-limited as one person", "Set TRUST_PROXY_HOPS"],
])
para(doc, "On a single server none of this applies, and the check stays silent.", grey=True)

doc.add_heading("A checklist", level=1)
doc.add_heading("Before going live", level=2)
bullets(doc, [
    "Two Super Admin accounts exist, and both passwords are known to different people.",
    "Every seeded demonstration account has been removed or had its password changed.",
    "Email is configured and a test message has arrived.",
    "WhatsApp has a token, a phone number id AND an approved template — or is deliberately left "
    "simulated.",
    "Google Drive is connected and each class folder is shared with the service account as Editor.",
    "Backups are running, and one has been restored somewhere to prove it works.",
    "The registration number format and the highest number already issued are correct.",
    "Institute policy has been reviewed: attendance thresholds, progress weights, what a "
    "certificate requires.",
    "The public page says what the Institute wants it to say, and the fee prices are published.",
    "TRUST_PROXY_HOPS is set if anything sits in front of the API.",
])
doc.add_heading("Every term", level=2)
bullets(doc, [
    "Create the new term and its intakes before sections are made.",
    "Check that every section has a teacher assigned.",
    "Review who still has an account and suspend those who have left.",
    "Read the audit log for role changes you did not expect.",
])

doc.add_page_break()
closing(doc, [
    "The startup log is the first place to look after any deployment change — it reports what is "
    "misconfigured before anybody notices from a screen.",
])

path = os.path.join(OUT, "Prepreneurship LMS - Super Administrator Guide.docx")
doc.save(path)
print("wrote", os.path.basename(path))
