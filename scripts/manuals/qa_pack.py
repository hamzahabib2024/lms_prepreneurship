"""
The QA pack — a document a tester can work from without reading any code.

WRITTEN FOR SOMEBODY WHO HAS NEVER SEEN THIS SYSTEM. That decides the shape:
credentials first, then how to tell a defect from a rule, then numbered cases
with explicit steps and an expected result for each.

THE SECTION THAT EARNS ITS PLACE is "things that look like bugs and are not".
Every one of these was reported as a defect during automated QA and turned out
to be the System behaving correctly. A tester who does not know them will
spend a day on each.
"""

import os
from docx import Document
from docx.shared import Pt
from build import (OUT, styled, cover, para, bullets, steps, callout, table)

doc = styled(Document())

cover(doc, "QA Test Pack",
      "How to test this System, and how to tell a defect from a rule",
      "For the QA engineer. No development experience assumed.")

# ───────────────────────────────────────────────────────────── contents ──
doc.add_heading("What is in this pack", level=1)
table(doc, ["Batch", "What it covers"], [
    ["1. Before you start", "What you need, and the accounts to use"],
    ["2. How to report", "What makes a report actionable"],
    ["3. Things that are NOT bugs", "Read this before raising anything"],
    ["4. Test cases", "94 numbered cases across 14 areas"],
    ["5. Automated checks", "What already runs, so you do not repeat it"],
    ["6. Sign-off", "What must pass before release"],
])
doc.add_page_break()

# ═══════════════════════════════════════════════════ 1. before you start ══
doc.add_heading("1. Before you start", level=1)

doc.add_heading("What you need", level=2)
bullets(doc, [
    "The web address of the System, from the Institute.",
    "A desktop browser (Chrome or Firefox) AND a phone. Several screens behave "
    "differently on a narrow screen and both must be checked.",
    "A microphone, for the voice recording cases.",
    "A small PDF, a JPEG, and one file renamed to the wrong extension — "
    "instructions for making that one are in case 5.7.",
])

doc.add_heading("The four roles", level=2)
para(doc, "Almost every defect in a system like this is a role seeing something "
          "they should not. Test each case as the role named, and where a case says "
          "'must not', check it as a DIFFERENT role too.")
table(doc, ["Role", "Sign-in", "Password"], [
    ["Super Admin", "superadmin@institute.local", "ChangeMe!SuperAdmin2026"],
    ["Admin (office)", "admin@institute.local", "ChangeMe!Admin2026"],
    ["Teacher", "sana@institute.local", "ChangeMe!Teacher2026"],
    ["Student", "hina2@student.local", "ChangeMe!Student2026"],
    ["Second student", "ayesha1@student.local", "ChangeMe!Student2026"],
])
callout(doc, "These are seeded demonstration accounts.",
        "They must be removed or have their passwords changed before the Institute goes live. "
        "If they still work on the production system, that is itself a defect — raise it.")

# ═══════════════════════════════════════════════════════ 2. how to report ══
doc.add_heading("2. How to report a defect", level=1)
para(doc, "A report that cannot be reproduced cannot be fixed. Include all five:")
steps(doc, [
    "WHO you were signed in as, by role.",
    "WHAT you did, as numbered steps from signing in.",
    "WHAT YOU EXPECTED, in one sentence.",
    "WHAT HAPPENED instead.",
    "THE REFERENCE, if an error appeared. It looks like ERR-A1B2C3D4 and identifies the exact "
    "request in the server log — with it, a developer can find the cause in minutes.",
])
callout(doc, "Severity, so triage is not guesswork:",
        "BLOCKER — data is lost, or one person can see another's records. "
        "MAJOR — a task cannot be completed at all. "
        "MINOR — a task is possible but awkward or wrongly worded. "
        "Anything where one role sees another's data is a BLOCKER regardless of how it looks.")

doc.add_page_break()

# ══════════════════════════════════════════ 3. things that are not bugs ═══
doc.add_heading("3. Things that look like bugs and are not", level=1)
para(doc, "Every one of these was reported during automated testing and turned out to be "
          "the System behaving correctly. Read them before raising anything.", bold=True)

table(doc, ["What you will see", "Why it is correct"], [
    ["A student opens the fee ledger and gets an empty list rather than 'access denied'.",
     "The permission says a student may read PAYMENTS — their own. The data layer then returns "
     "only their rows, and there are none in a list of debtors. Empty is the right answer. "
     "Raise it only if a student sees ANOTHER student's figures."],
    ["An admin opens the fee ledger and is asked for their password again.",
     "Anything touching money requires re-confirmation, even when already signed in. It means "
     "an unattended screen cannot be used to move money."],
    ["A teacher is refused the revenue report.",
     "Financial reporting is granted to an admin individually. A teacher never holds it."],
    ["Marks do not appear for a student after the teacher enters them.",
     "Marks are hidden until the teacher RELEASES the assignment, so a class sees them together. "
     "Check the teacher pressed Release."],
    ["A recording is listed but will not play.",
     "Check whether it says AVAILABLE. A recording whose file has been moved in Google Drive is "
     "marked MISSING, which is the System reporting a real problem, not causing one."],
    ["A student sees 'Anonymous' where a name should be.",
     "Students may post anonymously. Staff always see who it was. If a STUDENT can see the name "
     "of an anonymous poster, that IS a defect and a blocker."],
    ["A quiet week shows an empty calendar grid rather than a message.",
     "An empty month is the shape of that month. Only the list view says 'no classes'."],
    ["WhatsApp messages are not delivered.",
     "If the Institute has not set a token, a phone number id AND an approved template, messages "
     "are SIMULATED and readable on the Integrations screen. Check there before raising it."],
])

doc.add_page_break()

# ═════════════════════════════════════════════════════════ 4. test cases ══
doc.add_heading("4. Test cases", level=1)
para(doc, "Each case names the role, the steps, and the expected result. A case fails if the "
          "expected result does not happen exactly.", grey=True)


def area(n, title, note=None):
    doc.add_heading(f"{n}. {title}", level=2)
    if note:
        para(doc, note, grey=True)


def cases(rows):
    table(doc, ["#", "Role", "Do this", "Expect"], rows)


area(1, "Signing in and sessions")
cases([
    ["1.1", "Any", "Sign in with the correct password", "You reach the dashboard"],
    ["1.2", "Any", "Sign in with a wrong password", "Refused, and the reason does not say whether the EMAIL exists"],
    ["1.3", "Any", "Get the password wrong several times in a row", "The account locks temporarily and says so"],
    ["1.4", "Student", "Sign in, then reload the page", "Still signed in — a reload must not sign you out"],
    ["1.5", "Any", "Sign out, then press the browser Back button", "You are NOT returned to a signed-in screen"],
    ["1.6", "Any", "Change your password, then sign in with the old one", "Refused"],
])

area(2, "What a stranger can reach", "Sign out completely, or use a private window.")
cases([
    ["2.1", "Nobody", "Open the public page", "Loads: programmes, videos, photographs"],
    ["2.2", "Nobody", "Open /verify and enter a made-up certificate number", "Says no certificate matches — NOT an error page"],
    ["2.3", "Nobody", "Open /track and enter a made-up reference", "Says it cannot be found, politely"],
    ["2.4", "Nobody", "Try to open /settings, /users, /audit directly", "Refused, every time"],
    ["2.5", "Nobody", "Submit the application form", "Accepted, with a tracking reference"],
])

area(3, "Admissions")
cases([
    ["3.1", "Admin", "Open Admissions", "Applications listed, newest first"],
    ["3.2", "Admin", "Open one and view the payment slip", "The slip opens"],
    ["3.3", "Admin", "Approve an application", "A student, a registration number and a sign-in are created together"],
    ["3.4", "Admin", "Decline one", "They are told, and the reason is included"],
    ["3.5", "Teacher", "Try to open Admissions", "Not offered, and refused if the address is typed"],
])

area(4, "Attendance")
cases([
    ["4.1", "Teacher", "Take a register, marking one absent and one late", "Saves; everybody untouched is present"],
    ["4.2", "Teacher", "Reopen the register", "Your marks are still there"],
    ["4.3", "Teacher", "Change a mark and save", "Accepted, and the change is recorded"],
    ["4.4", "Student", "Check your own attendance", "Matches what the teacher recorded"],
    ["4.5", "Teacher", "Open a class you do NOT teach by typing its address", "Refused"],
])

area(5, "Assignments — the largest area")
cases([
    ["5.1", "Teacher", "Create an assignment and save it as a draft", "Created"],
    ["5.2", "Student", "Look for that assignment", "NOT visible — a draft reaches nobody"],
    ["5.3", "Teacher", "Publish it", "Now visible to the student"],
    ["5.4", "Teacher", "Choose 'a limited number of times', then set how many", "The number field appears and saves"],
    ["5.5", "Teacher", "Reopen the assignment", "The attempt limit and grace period you set are shown"],
    ["5.6", "Teacher", "Marking > an assignment > The brief > attach a PDF", "Attached and listed"],
    ["5.7", "Teacher", "Rename a JPEG to .pdf and attach it", "REFUSED — it says the contents are not a PDF"],
    ["5.8", "Teacher", "Attach the same file twice", "Accepted, not duplicated"],
    ["5.9", "Teacher", "Record a spoken brief", "Records, plays back, saves"],
    ["5.10", "Student", "Open the assignment", "Instructions, the attachment, and the spoken brief are all there"],
    ["5.11", "Student", "Download the attachment", "Downloads and opens"],
    ["5.12", "Student", "Look for a way to remove the teacher's file", "There is none — only Download is offered"],
    ["5.13", "Teacher", "Remove an attached file", "Gone, for the student too"],
    ["5.12", "Student", "Upload a file but do NOT press Submit, then reload", "Shows as NOT submitted — this is correct and important"],
    ["5.13", "Student", "Press Submit", "Confirmed on screen"],
    ["5.14", "Student", "Record a spoken answer and submit", "Appears in the list like any other file"],
    ["5.15", "Teacher", "Open Marking, choose the assignment, press Mark", "One submission at a time, with the mark box"],
    ["5.16", "Teacher", "Enter a mark ABOVE the maximum", "Refused"],
    ["5.17", "Teacher", "Enter a valid mark and save", "Moves to the next one still to mark"],
    ["5.18", "Teacher", "Press the right arrow key, then Escape", "Moves between submissions; Escape returns to the list"],
    ["5.19", "Student", "Check the mark before it is released", "NOT shown"],
    ["5.20", "Teacher", "Release the grades", "All students see theirs at once"],
    ["5.21", "Teacher", "Change a released mark", "A reason is required, and recorded"],
])

area(6, "Quizzes")
cases([
    ["6.1", "Teacher", "Build a quiz with a multiple-choice and a written question", "Saves"],
    ["6.2", "Teacher", "Publish it", "Visible to students when it opens"],
    ["6.3", "Student", "Start it, answer, submit", "Multiple choice marked immediately"],
    ["6.4", "Student", "Start a quiz and close the tab without submitting", "The attempt is used — check the rules said so"],
    ["6.5", "Teacher", "Open quiz marking", "Written answers grouped BY QUESTION, not by student"],
    ["6.6", "Teacher", "Mark and release", "Students see results"],
])

area(7, "Recordings and Google Drive")
cases([
    ["7.1", "Teacher", "Connect a Drive folder to a class", "Recordings from it appear by themselves"],
    ["7.2", "Teacher", "Play one", "Plays, and you can drag the scrubber to seek"],
    ["7.3", "Student", "Watch part of a lecture, leave, come back", "Resumes where you stopped"],
    ["7.4", "Student", "Drag to the end without watching", "Does NOT count as fully watched"],
    ["7.5", "Teacher", "Change the class's Drive folder", "Old recordings go, the new folder's arrive, and it says how many of each"],
    ["7.6", "Teacher", "Change it back", "The originals return with watch progress intact"],
])

area(8, "The class meeting link")
cases([
    ["8.1", "Teacher", "Courses > a class > Set the meeting link, beginning https://", "Saved, and a Join button appears"],
    ["8.2", "Teacher", "Set one beginning http:// (no s)", "REFUSED, and the message says it must begin https://"],
    ["8.3", "Teacher", "Add a note such as 'join five minutes early'", "Shown under the Join button"],
    ["8.4", "Student", "My subjects > the same subject", "A Join button, with the note, above the assignments"],
    ["8.5", "Student", "Look for any way to change the link", "There is none — no editor is offered at all"],
    ["8.6", "Teacher", "Clear the box and save", "The Join button disappears for the student too"],
    ["8.7", "Office", "Set a link on a class you do not teach", "Allowed — the office can set any"],
    ["8.8", "Office", "Batches > a batch > Subjects > the Meeting link column", "Each subject has its own link and its own teacher"],
    ["8.9", "Teacher", "Set the link from Batches rather than Courses", "Works — same editor, same rules"],
])

area(8.5, "Deleting a course, intake or subject")
cases([
    ["8.5.1", "Office", "Courses & fees > create a subject, then Delete it", "Gone from the list"],
    ["8.5.2", "Office", "Delete a subject a class is teaching", "REFUSED, naming the classes by code"],
    ["8.5.3", "Office", "Structure > Delete an intake that has batches", "REFUSED, naming the batches"],
    ["8.5.4", "Office", "Batches > Delete a batch that has subjects on it", "REFUSED, naming the subjects"],
    ["8.5.5", "Office", "Batches > open subjects > Remove one that has never been taught", "Removed"],
    ["8.5.6", "Office", "Remove a subject that HAS assignments or a register", "REFUSED, saying what it has"],
    ["8.5.7", "Teacher", "Look for any Delete on Batches", "There is none — no delete is offered at all"],
    ["8.5.8", "Office", "After any refusal, read the message", "It names what is in the way and what to do"],
])

area(9, "Discussion")
cases([
    ["9.1", "Student", "Ask a question normally", "Your name is shown"],
    ["9.2", "Student", "Ask one with 'Hide my name' ticked", "Shows as Anonymous"],
    ["9.3", "Second student", "Open that question", "Must show Anonymous — NOT the asker's name. A name here is a BLOCKER"],
    ["9.4", "Teacher", "Open the same question", "Shows the real name AND that it was anonymous"],
    ["9.5", "Teacher", "Mark a reply as the answer", "Marked, and stands out"],
    ["9.6", "Teacher", "Mark the question answered", "Shown as answered"],
    ["9.7", "Student", "Try to mark a question answered", "Refused"],
])

area(10, "Completion and certificates")
cases([
    ["10.1", "Teacher", "Open 'Who has finished' for a class", "Every student, with their figures"],
    ["10.2", "Teacher", "Mark somebody complete who does NOT meet the requirements, without a reason", "Refused — a reason is required"],
    ["10.3", "Teacher", "Do the same WITH a reason", "Accepted, and recorded as an override"],
    ["10.4", "Student", "Try to sign yourself off", "Refused"],
    ["10.5", "Admin", "Issue a certificate to a signed-off student", "Issued, with a number"],
    ["10.6", "Admin", "Try to issue to somebody not signed off", "Refused, naming what is outstanding"],
    ["10.7", "Nobody", "Verify that certificate number on the public page", "Confirmed genuine, with no account"],
    ["10.8", "Admin", "Revoke it, then verify again", "Now reads as revoked — NOT as missing"],
])

area(11, "Fees")
cases([
    ["11.1", "Student", "Open Fees", "What you owe, and each instalment"],
    ["11.2", "Student", "Submit a payment with a photo of a slip", "Shows as awaiting confirmation"],
    ["11.3", "Student", "Check the balance", "UNCHANGED until the office confirms — this is correct"],
    ["11.4", "Admin", "Confirm the payment", "Balance reduces and a receipt appears"],
    ["11.5", "Student", "Open the receipt", "Opens, and matches the amount"],
    ["11.6", "Teacher", "Try to open the fee ledger", "Refused"],
])

area(12, "Announcements")
cases([
    ["12.1", "Admin", "Post to 'Everyone at the Institute'", "Students and staff both see it"],
    ["12.2", "Admin", "Post to 'Teachers only'", "Teacher sees it"],
    ["12.3", "Student", "Look for that notice", "NOT visible. Visible is a BLOCKER"],
    ["12.4", "Admin", "Post to 'The public page only'", "Appears on the public website"],
    ["12.5", "Student", "Look for it in the System", "NOT visible — it was for visitors"],
    ["12.6", "Teacher", "Try to post to all teachers", "Refused — that is the office's"],
    ["12.7", "Admin", "Post one with an expiry in the near future, then wait", "It disappears by itself"],
])

area(13, "The public page editor")
cases([
    ["13.1", "Admin", "Change the headline and save", "Live on the public page immediately"],
    ["13.2", "Admin", "Upload a photograph", "Appears in the gallery"],
    ["13.3", "Admin", "Paste a YouTube link", "Appears as a playable card"],
    ["13.4", "Admin", "Press 'Restore the default' on a field", "Returns to the original wording"],
    ["13.5", "Teacher", "Try to open the public page editor", "Refused"],
])

area(14, "On a phone", "Open the System on a real phone, not a resized browser window.")
cases([
    ["14.1", "Student", "Sign in and open the dashboard", "Readable without sideways scrolling"],
    ["14.2", "Student", "Open the timetable", "Defaults to the list view"],
    ["14.3", "Student", "Submit an assignment with a photo", "Works, including taking the photo"],
    ["14.4", "Student", "Record a spoken answer", "Microphone permission is requested and recording works"],
    ["14.5", "Teacher", "Take a register", "Usable one-handed"],
    ["14.6", "Any", "Check every screen for text cut off at the edge", "Nothing is cut off"],
])

doc.add_page_break()

# ════════════════════════════════════════════════ 5. automated checks ═════
doc.add_heading("5. What is already checked automatically", level=1)
para(doc, "So you do not spend time repeating it. These run on every change and the build "
          "fails if any of them does.")
table(doc, ["What", "How many"], [
    ["Unit and integration tests", "1,500"],
    ["End-to-end checks against a running System", "86"],
    ["Static guards (every model policed, every route reachable, every screen routed, "
     "every icon and CSS class defined, accessibility basics)", "7 suites"],
])
para(doc, "What automation CANNOT check, and therefore matters most from you: whether a screen "
          "makes sense to somebody who has not seen it before; whether the wording is right; "
          "whether it works on a real phone on a slow connection; and whether the SYSTEM'S "
          "answer matches what the Institute actually intends.", bold=True)

# ═══════════════════════════════════════════════════════════ 6. sign-off ══
doc.add_heading("6. Before release", level=1)
para(doc, "Every one of these must pass. Any failure is a blocker.")
bullets(doc, [
    "No role can see another person's marks, fees, or personal details.",
    "A student cannot see the name behind an anonymous discussion post.",
    "A draft assignment or unpublished recording reaches no student.",
    "Marks are invisible until released.",
    "A certificate cannot be issued without a sign-off, and verifies publicly once issued.",
    "A payment does not reduce a balance until the office confirms it.",
    "Every seeded demonstration account has been removed or had its password changed.",
    "The System works on a phone.",
])

doc.add_paragraph()
para(doc, "If anything in this pack does not match what you see, the pack is wrong and that is "
          "worth reporting too — it is generated alongside the System and can drift.", grey=True)

path = os.path.join(OUT, "Prepreneurship LMS - QA Test Pack.docx")
doc.save(path)
print("wrote", os.path.basename(path))
